import io
import logging
import re

import math2docx
import requests
import supabase
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt
from fastapi import Depends, HTTPException, Request, Response
from pydantic import BaseModel

from api.v1.auth import get_supabase_client
from api.v1.qgen.paper_layout_config import (
    PAGE_MARGIN_BOTTOM_MM,
    PAGE_MARGIN_LEFT_MM,
    PAGE_MARGIN_RIGHT_MM,
    PAGE_MARGIN_TOP_MM,
)
from api.v1.qgen.utils.paper_utils import fetch_paper_data, format_duration

logger = logging.getLogger(__name__)


class DownloadDocxRequest(BaseModel):
    draft_id: str
    mode: str  # "paper" or "answer"


async def download_docx(
    download_req: DownloadDocxRequest,
    fastapi_request: Request,
    supabase_client: supabase.Client = Depends(get_supabase_client),
):
    """
    API endpoint to generate and download a DOCX of the question paper.
    """
    logger.info(
        "Received download_docx request",
        extra={"draft_id": download_req.draft_id, "mode": download_req.mode},
    )

    # 1. Fetch Paper Data
    data = await fetch_paper_data(download_req.draft_id, supabase_client)

    draft = data["draft"]
    sections = data["sections"]
    questions = data["questions"]
    instructions = data["instructions"]
    logo_url = data["logo_url"]
    images_map = data["images_map"]

    # Get toggle values from draft (default to True for backward compatibility)
    show_logo = draft.get("is_show_logo", True)
    show_instructions = draft.get("is_show_instruction", True)
    show_explanation = draft.get("is_show_explanation_answer_key", True)

    latex_regex = r"(\$\$.*?\$\$|\$.*?\$|\\\(.*?\\\)|\\\[.*?\\\])"

    def set_table_no_border(table):
        """Remove all borders from a table to make it invisible."""
        tbl = table._tbl
        # Get or create tblPr element
        tbl_pr = tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl.insert(0, tbl_pr)

        # Create and add borders element
        tbl_borders = OxmlElement("w:tblBorders")
        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "nil")
            tbl_borders.append(border)
        tbl_pr.append(tbl_borders)

    def add_horizontal_rule(paragraph):
        """Adds a bottom border to the paragraph to create a horizontal line."""
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")

        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")  # 6 = 3/4 pt
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        pBdr.append(bottom)

        # Ensure pBdr is inserted before alignment (jc) or other spacing elements
        # if they exist to respect some schema strictness, though mostly for jc.
        # Ideally pBdr is early in pPr.
        # Find first element that should come AFTER pBdr?
        # Schema: pStyle, keepNext, keepLines, pageBreakBefore, framePr,
        #         widowControl, numPr, suppressLineNumbers, pBdr
        # So we append, unless we see something that must be after.
        # But simply appending usually works for 'jc' if 'jc' isn't set yet.
        # If we set alignment later, python-docx handles it?
        # Let's just insert at the beginning of pPr to be safe if no pStyle,
        # or after pStyle if it exists.

        if len(pPr) > 0 and pPr[0].tag == qn("w:pStyle"):
            pPr.insert(1, pBdr)
            return

        pPr.insert(0, pBdr)  # Insert as first element/early element safe bet for fresh para

    def remove_control_characters(text):
        """
        Removes non-printable control characters that are invalid in XML 1.0.
        Valid characters are:
        #x9 | #xA | #xD | [#x20-#xD7FF] | [#xE000-#xFFFD] | [#x10000-#x10FFFF]
        """
        if not text:
            return ""
        return "".join(
            ch
            for ch in text
            if (
                (0x20 <= ord(ch) <= 0xD7FF)
                or (0xE000 <= ord(ch) <= 0xFFFD)
                or (0x10000 <= ord(ch) <= 0x10FFFF)
                or ch in ("\t", "\n", "\r")
            )
        )

    def add_text_with_math(paragraph, text, context_info=""):
        if not text:
            return
        parts = re.split(latex_regex, text)
        for part in parts:
            if not part:
                continue

            # Check if part is LaTeX
            is_latex = re.match(latex_regex, part)
            if is_latex:
                # Clean delimiters
                clean_latex = part
                if part.startswith("$$") and part.endswith("$$"):
                    clean_latex = part[2:-2]
                elif part.startswith("$") and part.endswith("$"):
                    clean_latex = part[1:-1]
                elif part.startswith("\\[") and part.endswith("\\]"):
                    clean_latex = part[2:-2]
                elif part.startswith("\\(") and part.endswith("\\)"):
                    clean_latex = part[2:-2]

                try:
                    # Capture state before attempt
                    initial_msg_count = len(paragraph._p)
                    # logger.info(
                    #     f"DEBUG: Before math '{clean_latex[:20]}', "
                    #     f"children: {initial_msg_count}"
                    # )

                    math2docx.add_math(paragraph, clean_latex)

                    # logger.info(f"DEBUG: Success adding math. Children: {len(paragraph._p)}")
                except Exception as e:
                    logger.warning(f"Failed to add math [{context_info}]: '{part}' -> {e}")

                    # Rollback
                    current_msg_count = len(paragraph._p)
                    # logger.debug(
                    #     f"DEBUG: Rollback triggered. Initial: {initial_msg_count}, "
                    #     f"Current: {current_msg_count}"
                    # )

                    if current_msg_count > initial_msg_count:
                        diff = current_msg_count - initial_msg_count
                        logger.debug(f"DEBUG: Removing {diff} elements causing corruption.")
                        for _ in range(diff):
                            # Remove the last element
                            if len(paragraph._p) > 0:
                                removed = paragraph._p[-1]
                                # logger.debug(f"DEBUG: Removing XML tag: {removed.tag}")
                                paragraph._p.remove(removed)

                    try:
                        sanitized_part = remove_control_characters(part)
                        paragraph.add_run(sanitized_part)  # Fallback to text
                    except Exception as e_run:
                        logger.error(f"Fallback add_run failed for part '{part}': {e_run}")
            else:
                try:
                    sanitized_part = remove_control_characters(part)
                    paragraph.add_run(sanitized_part)
                except Exception as e_run:
                    logger.error(f"Standard add_run failed for part '{part}': {e_run}")

    # 3. Build DOCX
    try:
        doc = Document()

        # Set page margins from config (synced with frontend/PDF)
        section = doc.sections[0]
        section.top_margin = Mm(PAGE_MARGIN_TOP_MM)
        section.right_margin = Mm(PAGE_MARGIN_RIGHT_MM)
        section.bottom_margin = Mm(PAGE_MARGIN_BOTTOM_MM)
        section.left_margin = Mm(PAGE_MARGIN_LEFT_MM)

        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(12)

        # Header Section ... (rest remains similar but using add_text_with_math)
        if show_logo and logo_url:
            try:
                response = requests.get(logo_url)
                if response.status_code == 200:
                    image_stream = io.BytesIO(response.content)
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(image_stream, height=Inches(0.6))
            except Exception as e:
                logger.warning(f"Failed to add logo to DOCX: {e}")

        # Institute Name
        h1 = doc.add_paragraph()
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = h1.add_run(draft.get("institute_name") or "Institute Name")
        run.bold = True
        run.font.size = Pt(20)

        # Paper Title
        h2 = doc.add_paragraph()
        h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_suffix = " - Answer Key" if download_req.mode == "answer" else ""
        run = h2.add_run((draft.get("paper_title") or "Examination Paper") + title_suffix)
        run.bold = True
        run.font.size = Pt(16)

        # Meta Information Table
        table = doc.add_table(rows=2, cols=2)
        table.width = Inches(6)
        set_table_no_border(table)  # Hide borders

        # Row 1
        cells_r1 = table.rows[0].cells
        cells_r1[0].text = f"Subject: {draft.get('subject_name') or '..........'}"
        cells_r1[1].text = f"Class: {draft.get('school_class_name') or '..........'}"
        cells_r1[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Row 2
        cells_r2 = table.rows[1].cells
        cells_r2[0].text = f"Max. Marks: {draft.get('maximum_marks') or '...'}"
        cells_r2[1].text = f"Duration: {format_duration(draft.get('paper_duration'))}"
        cells_r2[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Horizontal Rule
        hr_p = doc.add_paragraph()
        add_horizontal_rule(hr_p)

        # Instructions
        if download_req.mode == "paper" and show_instructions and instructions:
            doc.add_paragraph("General Instructions:").runs[0].bold = True
            for inst in instructions:
                p = doc.add_paragraph(style="List Number")
                add_text_with_math(p, inst.get("instruction_text"))

        # Calculate available width for right alignment
        doc_section = doc.sections[0]
        # Use safe defaults if margins are not set (though usually they are)
        page_width = doc_section.page_width or Inches(8.5)
        left_margin = doc_section.left_margin or Inches(1.0)
        right_margin = doc_section.right_margin or Inches(1.0)
        printable_width = page_width - left_margin - right_margin

        # Sections and Questions
        for section in sections:
            section_questions = sorted(
                [q for q in questions if q["qgen_draft_section_id"] == section["id"]],
                key=lambda q: q.get("position_in_draft", 0),
            )
            if not section_questions:
                continue

            total_marks = sum(q.get("marks", 0) for q in section_questions)

            p = doc.add_paragraph()
            p.paragraph_format.tab_stops.add_tab_stop(printable_width, WD_TAB_ALIGNMENT.RIGHT)

            run = p.add_run(f"{section.get('section_name')}")
            run.bold = True
            run.underline = True
            run.font.size = Pt(14)

            m_run = p.add_run(f"\t[{total_marks}]")
            m_run.bold = True
            m_run.font.size = Pt(14)

            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            for idx, q in enumerate(section_questions):
                q_p = doc.add_paragraph()

                # Add right-aligned tab stop
                q_p.paragraph_format.tab_stops.add_tab_stop(printable_width, WD_TAB_ALIGNMENT.RIGHT)

                q_p.add_run(f"{idx + 1}. ").bold = True

                # Question Text with Math
                add_text_with_math(q_p, q.get("question_text", ""), f"Q{idx + 1}-Text")

                m_run = q_p.add_run(f"\t[{q.get('marks')}]")
                m_run.italic = True
                q_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # Images
                q_images = images_map.get(q["id"], [])
                for img in q_images:
                    if img.get("img_url"):
                        try:
                            resp = requests.get(img["img_url"])
                            if resp.status_code == 200:
                                doc.add_picture(io.BytesIO(resp.content), width=Inches(2.5))
                        except Exception:
                            pass

                # Options (MCQ)
                if download_req.mode == "paper":
                    if q.get("question_type") in ["mcq4", "msq4"]:
                        opt_table = doc.add_table(rows=2, cols=2)
                        set_table_no_border(opt_table)  # Hide borders
                        opts = [
                            q.get("option1"),
                            q.get("option2"),
                            q.get("option3"),
                            q.get("option4"),
                        ]
                        labels = ["a) ", "b) ", "c) ", "d) "]
                        for i, opt in enumerate(opts):
                            row = i // 2
                            col = i % 2
                            if opt:
                                cell_p = opt_table.rows[row].cells[col].paragraphs[0]
                                cell_p.add_run(labels[i]).bold = True
                                add_text_with_math(cell_p, str(opt), f"Q{idx + 1}-Opt{i + 1}")
                    elif q.get("question_type") == "match_the_following":
                        cols = q.get("match_the_following_columns") or {}
                        col_names = list(cols.keys())
                        if len(col_names) >= 2:
                            left_col = cols[col_names[0]]
                            right_col = cols[col_names[1]]
                            max_rows = max(len(left_col), len(right_col))

                            match_table = doc.add_table(rows=max_rows + 1, cols=2)
                            set_table_no_border(match_table)

                            # Headers
                            for c_idx, name in enumerate(col_names[:2]):
                                cell_p = match_table.rows[0].cells[c_idx].paragraphs[0]
                                cell_p.add_run(name).bold = True

                            # Items
                            for r_idx in range(max_rows):
                                left_item = left_col[r_idx] if r_idx < len(left_col) else ""
                                right_item = right_col[r_idx] if r_idx < len(right_col) else ""

                                # Left Col
                                l_cell_p = match_table.rows[r_idx + 1].cells[0].paragraphs[0]
                                l_cell_p.add_run(f"{r_idx + 1}. ").bold = True
                                add_text_with_math(l_cell_p, str(left_item), f"Q{idx + 1}-L{r_idx}")

                                # Right Col
                                r_cell_p = match_table.rows[r_idx + 1].cells[1].paragraphs[0]
                                r_cell_p.add_run(f"{chr(65 + r_idx)}. ").bold = True
                                add_text_with_math(
                                    r_cell_p, str(right_item), f"Q{idx + 1}-R{r_idx}"
                                )

                # Answer Key
                if download_req.mode == "answer":
                    ans_p = doc.add_paragraph()
                    ans_p.add_run("Ans: ").bold = True
                    add_text_with_math(ans_p, str(q.get("answer_text") or "N/A"), f"Q{idx + 1}-Ans")

                    if show_explanation and q.get("explanation"):
                        exp_p = doc.add_paragraph()
                        exp_p.add_run("Explanation: ").bold = True
                        add_text_with_math(exp_p, str(q["explanation"]), f"Q{idx + 1}-Expl")

                # Page Break
                if q.get("is_page_break_below"):
                    doc.add_page_break()

        # Save to BytesIO
        target_stream = io.BytesIO()
        try:
            doc.save(target_stream)
        except Exception as e:
            logger.error(f"Failed to save DOCX (content corruption check): {e}")
            raise HTTPException(status_code=500, detail="Document generation corrupted") from e

        target_stream.seek(0)

        filename = f"{draft.get('paper_title', 'Paper')}_{download_req.mode}.docx"

        return Response(
            content=target_stream.read(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )

    except HTTPException:
        # Re-raise HTTPExceptions so they aren't caught by the general Exception block
        raise
    except Exception as e:
        logger.exception("DOCX generation failed")
        raise HTTPException(status_code=500, detail="Failed to generate DOCX") from e
